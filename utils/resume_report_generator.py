
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_resume_report(
    candidate_name,
    final_score,
    optimized_score,
    optimized,
    validation,
    recruiter,
    output_path
):

    doc = Document()

    # --------------------------------------------------
    # Title
    # --------------------------------------------------

    title = doc.add_heading(
        "AI Resume Optimizer Report",
        level=1
    )

    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --------------------------------------------------
    # Candidate Information
    # --------------------------------------------------

    doc.add_heading(
        "Candidate Information",
        level=2
    )

    doc.add_paragraph(
        f"Candidate Name : {candidate_name}"
    )

    # --------------------------------------------------
    # Executive Summary
    # --------------------------------------------------

    doc.add_heading(
        "Executive Summary",
        level=2
    )

    original_final = float(final_score)
    optimized_final = float(optimized_score["final_score"])

    improvement = round(
        optimized_final - original_final,
        2
    )

    doc.add_paragraph(
        f"Original ATS Score : {original_final:.2f}%"
    )

    doc.add_paragraph(
        f"Optimized ATS Score : {optimized_final:.2f}%"
    )

    doc.add_paragraph(
        f"ATS Improvement : +{improvement:.2f}%"
    )

    # --------------------------------------------------
    # Resume Validation
    # --------------------------------------------------

    doc.add_heading(
        "Resume Validation",
        level=2
    )

    doc.add_paragraph(
        f"Truthfulness Score : {validation['truthfulness_score']}%"
    )

    doc.add_paragraph(
        f"Hallucination : {validation['hallucination_detected']}"
    )

    doc.add_paragraph(
        validation["summary"]
    )

    # --------------------------------------------------
    # Verified Sections
    # --------------------------------------------------

    doc.add_heading(
        "Verified Sections",
        level=2
    )

    for section in validation.get(
        "verified_sections",
        []
    ):
        doc.add_paragraph(
            f"• {section}"
        )

    # --------------------------------------------------
    # Warnings
    # --------------------------------------------------

    if validation.get("warnings"):

        doc.add_heading(
            "Warnings",
            level=2
        )

        for warning in validation["warnings"]:

            doc.add_paragraph(
                f"• {warning}"
            )

    # --------------------------------------------------
    # Optimization Summary
    # --------------------------------------------------

    doc.add_heading(
        "Optimization Summary",
        level=2
    )

    doc.add_paragraph(
        optimized.get("summary", "")
    )

    # --------------------------------------------------
    # Keywords Added
    # --------------------------------------------------

    doc.add_heading(
        "JD Keywords Added",
        level=2
    )

    for keyword in optimized.get(
        "keywords_added",
        []
    ):

        doc.add_paragraph(
            f"• {keyword}"
        )

    # --------------------------------------------------
    # Keyword Mapping
    # --------------------------------------------------

    doc.add_heading(
        "Keyword Mapping",
        level=2
    )

    table = doc.add_table(
        rows=1,
        cols=2
    )

    table.style = "Table Grid"

    header = table.rows[0].cells

    header[0].text = "JD Keyword"
    header[1].text = "Resume Change"

    for item in optimized.get(
        "keyword_mapping",
        []
    ):

        row = table.add_row().cells

        row[0].text = item["jd_keyword"]
        row[1].text = item["resume_change"]

    # --------------------------------------------------
    # Resume Changes
    # --------------------------------------------------

    doc.add_heading(
        "Resume Changes",
        level=2
    )

    for change in optimized.get(
        "changes",
        []
    ):

        doc.add_paragraph(
            f"• {change}"
        )

    # --------------------------------------------------
    # ATS Improvements
    # --------------------------------------------------

    doc.add_heading(
        "ATS Improvements",
        level=2
    )

    for item in optimized.get(
        "ats_improvements",
        []
    ):

        doc.add_paragraph(
            f"• {item}"
        )

    # --------------------------------------------------
    # Compliance Report
    # --------------------------------------------------

    doc.add_heading(
        "Compliance Report",
        level=2
    )

    compliance = optimized["compliance"]

    doc.add_paragraph(
        f"Fake Skills : {compliance['fake_skills']}"
    )

    doc.add_paragraph(
        f"Fake Projects : {compliance['fake_projects']}"
    )

    doc.add_paragraph(
        f"Fake Experience : {compliance['fake_experience']}"
    )

    doc.add_paragraph(
        f"Fake Certifications : {compliance['fake_certifications']}"
    )

    doc.add_paragraph(
        f"ATS Safe : {compliance['ats_safe']}"
    )

    # --------------------------------------------------
    # Recruiter Verdict
    # --------------------------------------------------

    doc.add_heading(
        "Recruiter Verdict",
        level=2
    )

    doc.add_paragraph(
        recruiter["decision"]
    )

    doc.add_paragraph(
        recruiter["recommendation"]
    )

    doc.add_paragraph(
        recruiter["notes"]
    )

    # --------------------------------------------------
    # Save Report
    # --------------------------------------------------

    doc.save(output_path)