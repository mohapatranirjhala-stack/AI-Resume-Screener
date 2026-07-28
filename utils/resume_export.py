
from docx import Document


def export_resume_docx(resume_text, filename):

    doc = Document()

    doc.add_heading("ATS Optimized Resume", level=1)

    for line in resume_text.split("\n"):
        doc.add_paragraph(line)

    doc.save(filename)

    return filename